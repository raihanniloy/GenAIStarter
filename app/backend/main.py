import os
from typing import List

import PyPDF2
import openai
import supabase
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# Load environment variables
load_dotenv()

# Load API keys from environment variables
openai.api_key = os.getenv("OPENAI_API_KEY")
supabase_url = os.getenv("SUPABASE_URL")
supabase_key = os.getenv("SUPABASE_KEY")

# Initialize Supabase client
supabase_client = supabase.create_client(supabase_url, supabase_key)

app = FastAPI(title="AI-Powered Documentation Search API")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Next.js frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def generate_embedding(text: str):
    """Generate embedding using OpenAI API"""
    response = openai.Embedding.create(input=text, model="text-embedding-ada-002")
    return response["data"][0]["embedding"]


def extract_text_from_file(file: UploadFile) -> str:
    """Extract text from uploaded file based on its type"""
    try:
        if file.filename.endswith(".pdf"):
            reader = PyPDF2.PdfReader(file.file)
            text = "".join(page.extract_text() or "" for page in reader.pages)
        elif file.filename.endswith(".docx"):
            doc = Document(file.file)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif file.filename.endswith(".txt"):
            text = file.file.read().decode("utf-8")
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")


@app.post("/embed-multiple/")
async def embed_multiple_documents(files: List[UploadFile] = File(...)):
    """Upload and embed multiple documents"""
    results = []

    for file in files:
        text = extract_text_from_file(file)
        embedding = generate_embedding(text)

        response = supabase_client.table("documentation").insert({
            "title": file.filename,
            "content": text,
            "embedding": embedding
        }).execute()

        if response.get("error"):
            raise HTTPException(status_code=500, detail=f"Failed to store document: {file.filename}")

        results.append({"filename": file.filename, "status": "stored"})

    return {"message": "Documents processed successfully", "results": results}


class Document(BaseModel):
    title: str
    content: str


class SearchQuery(BaseModel):
    query: str


@app.post("/embed/")
def embed_document(doc: Document):
    """Embed a document and store it in Supabase"""
    embedding = generate_embedding(doc.content)

    response = supabase_client.table("documentation").insert({
        "title": doc.title,
        "content": doc.content,
        "embedding": embedding
    }).execute()

    if response.get("error"):
        raise HTTPException(status_code=500, detail="Failed to store document")

    return {"message": "Document stored successfully"}


@app.post("/search/")
def search_docs(query: SearchQuery):
    """Search documentation using vector similarity"""
    query_embedding = generate_embedding(query.query)

    # Perform similarity search using pgvector
    response = supabase_client.rpc("match_documents", {
        "query_embedding": query_embedding,
        "match_threshold": 0.7,
        "match_count": 5
    }).execute()

    if response.get("error"):
        raise HTTPException(status_code=500, detail="Search failed")

    return {"results": response["data"]}